'''
对于操作docx文件：
使用库：python-docx、pandas、pywin32
由于python-docx只能操作docx文件，所以必须将所操纵的文件doc改为docx，使用pywin32进行转换
'''
from win32com import client as wc
from docx import Document
import os
import re
import pdfplumber


class DocxOperate():

    def doctodocx(self, path):
        '''
        将doc文件批量转换为docx文件
        :param path:
        :return:
        '''
        # doc_list = [os.path.join(path, item) for item in os.listdir(path) if item.endswith('doc') or item.endswith('docx')]

        word = wc.Dispatch('Word.Application')
        # print(doc_list)
        for doc in doc_list:
            new_doc = doc.replace('doc', 'docx') if doc.endswith('doc') else doc
            operate_doc = word.Documents.Open(doc)
            operate_doc.SaveAs(new_doc, 12, False, '', True, '', False, False, False, False)
            operate_doc.Close()
            print('{} Save sucessfully '.format(new_doc))
            os.remove(doc) if doc.endswith('doc') else None
        word.Quit()

    # doctodocx('D:\program\pycharm\文件')

    def getText(self, path):
        '''
        读取所有段落的文字,读取书引号中的英文有问题
        :param path:
        :return:
        '''
        document = Document(path)
        all_paragraphs = document.paragraphs
        for item in all_paragraphs:
            print(item.text)

    # getText(r'../文件/1.docx')

    def getTable(self, path):
        '''
        对表格进行读取,数据提取有误差
        :param path:
        :return:
        '''
        # self.doctodocx(path)
        document = Document(path)
        # print(path)
        all_tables = document.tables
        item_list = []
        text_list = {}
        for table in all_tables:
            for row_index, row in enumerate(table.rows, 1):
                row_list = []
                for col_index, cell in enumerate(row.cells, 1):
                    print(col_index, cell.text)
                    item_list.append(cell.text) if row_index == 1 else row_list.append(cell.text)
                if row_list:
                    text_list[str(row_index - 1)] = row_list
                # print('第{}行第{}列的数据是-----{}'.format(row_index, col_index, cell.text))
        return item_list, text_list

    # item_list, text_list = getTable(r'../文件/zh.docx')
    # print(item_list)
    # print(text_list)

    def getPicture(self, path, result_path):
        '''
        进行docx中图片的提取，如果图片重复，只读取一张或者覆盖，如果存放地址有此文件，将会覆盖
        :param path:文档路径
        :param result_path:图片存放路径
        :return:
        '''
        document = Document(path)
        dict_rel = document.part._rels
        for rel in dict_rel:
            rel = dict_rel[rel]
            # print(type(rel))
            if 'image' in rel.target_ref:
                if not os.path.exists(result_path):
                    os.makedirs(result_path)
                # 获取文件中所有图片，并进行命名image2.png、image1.png......
                img_name = re.findall('/(.*)', rel.target_ref)[0]
                # 获取文件前缀名
                word_name = os.path.splitext(path)[0]
                if os.sep in word_name:
                    new_name = word_name.split('\\')[-1]
                else:
                    new_name = word_name.split('/')[-1]
                # 将文件名与图片名相结合
                img_name = f'{new_name}_{img_name}'
                # 将图片写入将要存放的位置
                with open(os.path.join(result_path, img_name), 'wb') as f:
                    # print(type(rel.target_part.blob))   #<class 'bytes'>
                    f.write(rel.target_part.blob)


# DocxOperate().getPicture(r'../文件/zh.docx', r'../文件')
# _, text_list = DocxOperate().getTable(r"C:\Users\hp\Desktop\2022大创\罗老师\计算机设计\118-字母词分级规范研究-王秋萍.docx")
# print(text_list)



class PDFOperate():
    def getTable(self, path):
        with pdfplumber.open(r'C:\Users\hp\Desktop\2022大创\罗老师\计算机设计\118-字母词分级规范研究-王秋萍.pdf') as pdf:
            print(pdf.pages)
            for i in range(16):
                page = pdf.pages[i]
                print("第" + str(i) + "页")
                for rows in page.extract_tables():
                    for row in rows:
                        row = list(filter(None, row))
                        print(row)
                    # print(row[0])  # 打印每个列表对应的第一个元素


def getDataFromTable(path):
    document = Document(path)
    # 获取文档中的所有表格
    all_tables = document.tables
    text = ''
    for table in all_tables:
        cell_list = []
        for row in table.rows:
            cell_list.clear()
            for cell in row.cells:
                tempt = text
                text = cell.text.replace('\n', '').replace(' ', '')
                if tempt == text:
                    continue
                cell_list.append(text)
            print(cell_list)


getDataFromTable(r'static\1.docx')


# PDFOperate().getTable(r'C:\Users\hp\Desktop\2022大创\罗老师\计算机设计\118-字母词分级规范研究-王秋萍.pdf')


# DocxOperate().getTable(r'static\1.docx')
